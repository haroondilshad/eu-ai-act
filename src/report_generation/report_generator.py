"""
Report generation module for EU AI Act compliance analysis results.
"""

import json
import logging
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Union, Any

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

from src.config import (
    OUTPUT_DIR,
    REPORT_TEMPLATE_PATH
)

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    stream=sys.stderr
)
logger = logging.getLogger("report_generator")

class ReportGenerator:
    """
    Generate compliance reports from analysis results.
    """
    
    def __init__(self, output_dir: Path = OUTPUT_DIR):
        """
        Initialize the report generator.
        
        Args:
            output_dir: Directory to store generated reports
        """
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_report(self, analysis_results: Union[Dict, Path, str]) -> Dict[str, Path]:
        """
        Generate a compliance report from analysis results.
        
        Args:
            analysis_results: Analysis results dictionary, or path to JSON file containing results
            
        Returns:
            Dictionary containing paths to the generated reports (docx and txt)
        """
        # Load analysis results if provided as path
        if isinstance(analysis_results, (str, Path)):
            with open(analysis_results, "r", encoding="utf-8") as f:
                results = json.load(f)
        else:
            results = analysis_results
        
        system_name = results.get("system_name", "Unnamed AI System")
        logger.info(f"Generating compliance report for {system_name}")
        
        # Create document
        doc = Document()
        
        # Add report title
        self._add_title(doc, f"EU AI Act Compliance Analysis: {system_name}")
        
        # Add executive summary
        self._add_executive_summary(doc, results)
        
        # Add compliance score visualization
        chart_path = self._generate_score_chart(results)
        self._add_image(doc, chart_path, width=6, caption="Compliance Scores by Category")
        
        # Add detailed analysis
        self._add_detailed_analysis(doc, results)
        
        # Add compliance gaps
        self._add_compliance_gaps(doc, results)
        
        # Add recommendations
        self._add_recommendations(doc, results)
        
        # Save report
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_filename = f"{system_name.replace(' ', '_')}_compliance_report_{timestamp}"
        docx_report_path = self.output_dir / f"{report_filename}.docx"
        doc.save(docx_report_path)
        
        # Generate and save text version of the report
        text_report_path = self._generate_text_report(results, report_filename)
        
        logger.info(f"Saved compliance report to {docx_report_path}")
        logger.info(f"Saved text version of report to {text_report_path}")
        
        return {
            "docx": docx_report_path,
            "txt": text_report_path
        }
    
    def _add_title(self, doc: Document, title: str) -> None:
        """
        Add title to the document.
        
        Args:
            doc: Document object
            title: Title text
        """
        doc.add_heading(title, level=0)
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_paragraph()  # Add spacing
    
    def _add_executive_summary(self, doc: Document, results: Dict) -> None:
        """
        Add executive summary to the report.
        
        Args:
            doc: Document object
            results: Analysis results
        """
        doc.add_heading("Executive Summary", level=1)
        
        system_type = results.get("system_type", "Undetermined")
        overall_score = results.get("overall_score", 0.0)
        
        # Format score as percentage
        score_percentage = f"{overall_score * 100:.1f}%"
        
        # Determine compliance level based on score and system type
        if system_type == "prohibited":
            compliance_level = "Prohibited Use"
            color = RGBColor(255, 0, 0)  # Red
        elif overall_score >= 0.8:
            compliance_level = "Highly Compliant"
            color = RGBColor(0, 128, 0)  # Green
        elif overall_score >= 0.6:
            compliance_level = "Moderately Compliant"
            color = RGBColor(255, 165, 0)  # Orange
        else:
            compliance_level = "Significant Gaps"
            color = RGBColor(255, 0, 0)  # Red
        
        # Summary paragraph
        summary = doc.add_paragraph()
        summary.add_run(f"System Type: ").bold = True
        summary.add_run(f"{system_type.capitalize()} AI System\n")
        
        summary.add_run(f"Overall Compliance Score: ").bold = True
        score_run = summary.add_run(f"{score_percentage} ")
        score_run.bold = True
        score_run.font.color.rgb = color
        
        summary.add_run(f"({compliance_level})\n\n")
        
        # Add summary based on system type
        if system_type == "prohibited":
            summary.add_run(
                "This AI system falls under the PROHIBITED category in the EU AI Act. "
                "Systems in this category cannot be legally deployed within the EU. "
                "See the detailed analysis section for specific prohibitions that apply and necessary changes."
            )
        elif system_type == "high-risk":
            summary.add_run(
                "As a high-risk AI system, this application is subject to the strictest requirements "
                "under the EU AI Act, including risk assessment, data governance, technical documentation, "
                "human oversight, and transparency obligations."
            )
        elif system_type == "general-purpose":
            summary.add_run(
                "As a general-purpose AI system, this application must comply with transparency "
                "requirements, copyright obligations, and certain technical documentation standards "
                "under the EU AI Act."
            )
        elif system_type == "limited-risk":
            summary.add_run(
                "As a limited-risk AI system with specific transparency obligations, this application "
                "must ensure users are aware when interacting with AI, and must comply with "
                "specific labeling and transparency requirements."
            )
        else:
            summary.add_run(
                "As a minimal-risk AI system, this application has limited obligations under the EU AI Act, "
                "but should still maintain appropriate documentation and risk management practices."
            )
        
        doc.add_paragraph()  # Add spacing
    
    def _generate_score_chart(self, results: Dict) -> Path:
        """
        Generate a radar chart of compliance scores by category.
        
        Args:
            results: Analysis results
            
        Returns:
            Path to the generated chart image
        """
        category_scores = results.get("category_scores", {})
        
        if not category_scores:
            logger.warning("No category scores found for chart generation")
            return None
        
        # Prepare data for radar chart
        categories = list(category_scores.keys())
        scores = [category_scores[cat] for cat in categories]
        
        # Format category names for display
        display_categories = [cat.replace("_", " ").title() for cat in categories]
        
        # Create the radar chart
        fig, ax = plt.subplots(figsize=(8, 6), subplot_kw=dict(polar=True))
        
        # Number of categories
        N = len(categories)
        
        # Angle of each axis
        angles = [n / float(N) * 2 * np.pi for n in range(N)]
        angles += angles[:1]  # Close the loop
        
        # Add the first point at the end to close the loop
        scores_plot = scores + [scores[0]]
        
        # Draw the chart
        ax.plot(angles, scores_plot, linewidth=2, linestyle='solid')
        ax.fill(angles, scores_plot, alpha=0.25)
        
        # Set category labels
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(display_categories)
        
        # Set y-axis limits
        ax.set_ylim(0, 1)
        ax.set_yticks([0.2, 0.4, 0.6, 0.8, 1.0])
        ax.set_yticklabels(['20%', '40%', '60%', '80%', '100%'])
        
        # Add grid
        ax.grid(True)
        
        # Add title
        plt.title('Compliance Scores by Category', size=14, y=1.1)
        
        # Save chart
        chart_path = self.output_dir / f"compliance_chart_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
        plt.tight_layout()
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        logger.info(f"Generated compliance score chart at {chart_path}")
        return chart_path
    
    def _add_image(self, doc: Document, image_path: Path, width: float = 6, caption: str = None) -> None:
        """
        Add an image to the document.
        
        Args:
            doc: Document object
            image_path: Path to the image
            width: Width of the image in inches
            caption: Optional caption for the image
        """
        if not image_path or not image_path.exists():
            logger.warning(f"Image not found: {image_path}")
            return
        
        # Add the image
        doc.add_picture(str(image_path), width=Inches(width))
        
        # Add caption if provided
        if caption:
            caption_paragraph = doc.add_paragraph(caption)
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_paragraph.style = 'Caption'
        
        doc.add_paragraph()  # Add spacing
    
    def _add_detailed_analysis(self, doc: Document, results: Dict) -> None:
        """
        Add detailed analysis section to the report.
        
        Args:
            doc: Document object
            results: Analysis results
        """
        doc.add_heading("Detailed Analysis", level=1)
        
        system_type = results.get("system_type", "undetermined")
        
        # For prohibited systems, add prohibition analysis
        if system_type == "prohibited":
            self._add_prohibition_analysis(doc, results)
            return
        
        detailed_analysis = results.get("detailed_analysis", {})
        
        if not detailed_analysis:
            doc.add_paragraph("No detailed analysis available.")
            return
        
        for category, analysis in detailed_analysis.items():
            # Format category name for display
            display_category = category.replace("_", " ").title()
            
            # Add category heading
            doc.add_heading(f"{display_category}", level=2)
            
            # Add score
            score = results.get("category_scores", {}).get(category, 0.0)
            score_percentage = f"{score * 100:.1f}%"
            
            score_paragraph = doc.add_paragraph()
            score_paragraph.add_run(f"Compliance Score: ").bold = True
            score_paragraph.add_run(f"{score_percentage}")
            
            # Add summary if available
            if "summary" in analysis:
                summary_paragraph = doc.add_paragraph()
                summary_paragraph.add_run("Summary: ").bold = True
                summary_paragraph.add_run(analysis["summary"])
            
            # Add findings table
            if "findings" in analysis and analysis["findings"]:
                doc.add_paragraph()
                doc.add_heading("Findings", level=3)
                
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = "Requirement"
                header_cells[1].text = "Compliance Level"
                header_cells[2].text = "Evidence / Gap"
                
                # Make header bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for finding in analysis["findings"]:
                    row_cells = table.add_row().cells
                    row_cells[0].text = finding.get("requirement", "")
                    row_cells[1].text = finding.get("compliance_level", "").replace("_", " ").title()
                    
                    # Combine evidence and gap
                    evidence = finding.get("evidence", "")
                    gap = finding.get("gap", "")
                    if evidence and gap:
                        row_cells[2].text = f"Evidence: {evidence}\n\nGap: {gap}"
                    elif evidence:
                        row_cells[2].text = evidence
                    elif gap:
                        row_cells[2].text = f"Gap: {gap}"
            
            doc.add_paragraph()  # Add spacing
    
    def _add_prohibition_analysis(self, doc: Document, results: Dict) -> None:
        """
        Add prohibition analysis for prohibited systems.
        
        Args:
            doc: Document object
            results: Analysis results
        """
        doc.add_heading("Prohibition Analysis", level=2)
        
        prohibition_analysis = results.get("prohibition_analysis", {})
        
        if not prohibition_analysis:
            doc.add_paragraph("No detailed prohibition analysis available.")
            return
        
        # Add summary paragraph
        summary = prohibition_analysis.get("summary", "")
        if summary:
            summary_paragraph = doc.add_paragraph()
            summary_paragraph.add_run("Summary: ").bold = True
            summary_paragraph.add_run(summary)
            
            doc.add_paragraph()  # Add spacing
        
        # Add social scoring analysis
        if "social_scoring" in prohibition_analysis:
            doc.add_heading("Social Scoring Analysis (Article 5(1)(c))", level=3)
            
            social_scoring = prohibition_analysis["social_scoring"]
            is_applicable = social_scoring.get("is_applicable", "no")
            analysis_text = social_scoring.get("analysis", "")
            
            applicable_paragraph = doc.add_paragraph()
            applicable_paragraph.add_run("Applicable: ").bold = True
            applicable_run = applicable_paragraph.add_run(is_applicable.upper())
            applicable_run.bold = True
            
            if is_applicable.lower() == "yes":
                applicable_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            
            analysis_paragraph = doc.add_paragraph()
            analysis_paragraph.add_run("Analysis: ").bold = True
            analysis_paragraph.add_run(analysis_text)
            
            doc.add_paragraph()  # Add spacing
        
        # Add manipulation techniques analysis
        if "manipulation_techniques" in prohibition_analysis:
            doc.add_heading("Manipulation Techniques Analysis (Article 5(1)(a))", level=3)
            
            manipulation = prohibition_analysis["manipulation_techniques"]
            is_applicable = manipulation.get("is_applicable", "no")
            analysis_text = manipulation.get("analysis", "")
            
            applicable_paragraph = doc.add_paragraph()
            applicable_paragraph.add_run("Applicable: ").bold = True
            applicable_run = applicable_paragraph.add_run(is_applicable.upper())
            applicable_run.bold = True
            
            if is_applicable.lower() == "yes":
                applicable_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            
            analysis_paragraph = doc.add_paragraph()
            analysis_paragraph.add_run("Analysis: ").bold = True
            analysis_paragraph.add_run(analysis_text)
            
            doc.add_paragraph()  # Add spacing
        
        # Add warning box
        warning = doc.add_paragraph()
        warning.paragraph_format.left_indent = Inches(0.5)
        warning.paragraph_format.right_indent = Inches(0.5)
        warning_run = warning.add_run(
            "⚠️ WARNING: This AI system falls under prohibited uses in the EU AI Act. "
            "It cannot be legally deployed in the EU without significant redesign to remove the prohibited elements."
        )
        warning_run.bold = True
        warning_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
    
    def _add_compliance_gaps(self, doc: Document, results: Dict) -> None:
        """
        Add compliance gaps section to the report.
        
        Args:
            doc: Document object
            results: Analysis results
        """
        doc.add_heading("Compliance Gaps", level=1)
        
        gaps = results.get("compliance_gaps", [])
        
        if not gaps:
            doc.add_paragraph("No significant compliance gaps identified.")
            return
        
        # Add introduction
        doc.add_paragraph(
            "The following compliance gaps have been identified and should be addressed "
            "to improve compliance with the EU AI Act requirements:"
        )
        
        # Add table of gaps
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Category"
        header_cells[1].text = "Requirement"
        header_cells[2].text = "Gap"
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        for gap in gaps:
            row_cells = table.add_row().cells
            row_cells[0].text = gap.get("category", "").replace("_", " ").title()
            row_cells[1].text = gap.get("requirement", "")
            row_cells[2].text = gap.get("gap", "")
        
        doc.add_paragraph()  # Add spacing
    
    def _add_recommendations(self, doc: Document, results: Dict) -> None:
        """
        Add recommendations section to the report.
        
        Args:
            doc: Document object
            results: Analysis results
        """
        doc.add_heading("Recommendations", level=1)
        
        recommendations = results.get("recommendations", [])
        
        if not recommendations:
            doc.add_paragraph("No specific recommendations available.")
            return
        
        # Add introduction
        doc.add_paragraph(
            "Based on the compliance analysis, the following recommendations are provided to "
            "improve compliance with the EU AI Act:"
        )
        
        # Sort recommendations by priority
        priority_order = {"high": 0, "medium": 1, "low": 2}
        sorted_recommendations = sorted(
            recommendations,
            key=lambda x: priority_order.get(x.get("priority", "low").lower(), 3)
        )
        
        # Add recommendations
        for i, rec in enumerate(sorted_recommendations, 1):
            category = rec.get("category", "").replace("_", " ").title()
            priority = rec.get("priority", "").upper()
            recommendation = rec.get("recommendation", "")
            
            # Format recommendation
            rec_paragraph = doc.add_paragraph(style='List Number')
            rec_paragraph.add_run(f"{category} ").bold = True
            
            # Add priority with color coding
            priority_run = rec_paragraph.add_run(f"({priority}) ")
            priority_run.bold = True
            
            if priority.lower() == "high":
                priority_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            elif priority.lower() == "medium":
                priority_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
            else:
                priority_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            
            rec_paragraph.add_run(recommendation)
        
        doc.add_paragraph()  # Add spacing
    
    def _generate_text_report(self, results: Dict, report_filename: str) -> Path:
        """
        Generate a text version of the compliance report.
        
        Args:
            results: Analysis results
            report_filename: Base filename for the report
            
        Returns:
            Path to the generated text report
        """
        system_name = results.get("system_name", "Unnamed AI System")
        system_type = results.get("system_type", "Undetermined")
        overall_score = results.get("overall_score", 0.0)
        
        # Format score as percentage
        score_percentage = f"{overall_score * 100:.1f}%"
        
        # Determine compliance level based on score and system type
        if system_type == "prohibited":
            compliance_level = "Prohibited Use"
        elif overall_score >= 0.8:
            compliance_level = "Highly Compliant"
        elif overall_score >= 0.6:
            compliance_level = "Moderately Compliant"
        else:
            compliance_level = "Significant Gaps"
        
        # Build text report content
        report_lines = [
            f"EU AI Act Compliance Analysis: {system_name}",
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            "EXECUTIVE SUMMARY",
            "=================",
            f"System Type: {system_type.capitalize()} AI System",
            f"Overall Compliance Score: {score_percentage} ({compliance_level})",
            ""
        ]
        
        # Add summary based on system type
        if system_type == "prohibited":
            report_lines.append(
                "This AI system falls under the PROHIBITED category in the EU AI Act. "
                "Systems in this category cannot be legally deployed within the EU. "
                "See the detailed analysis section for specific prohibitions that apply and necessary changes."
            )
        elif system_type == "high-risk":
            report_lines.append(
                "As a high-risk AI system, this application is subject to the strictest requirements "
                "under the EU AI Act, including risk assessment, data governance, technical documentation, "
                "human oversight, and transparency obligations."
            )
        elif system_type == "general-purpose":
            report_lines.append(
                "As a general-purpose AI system, this application must comply with transparency "
                "requirements, copyright obligations, and certain technical documentation standards "
                "under the EU AI Act."
            )
        elif system_type == "limited-risk":
            report_lines.append(
                "As a limited-risk AI system with specific transparency obligations, this application "
                "must ensure users are aware when interacting with AI, and must comply with "
                "specific labeling and transparency requirements."
            )
        else:
            report_lines.append(
                "As a minimal-risk AI system, this application has limited obligations under the EU AI Act, "
                "but should still maintain appropriate documentation and risk management practices."
            )
        
        report_lines.extend([
            "",
            "COMPLIANCE SCORES BY CATEGORY",
            "=============================",
        ])
        
        # Add category scores
        category_scores = results.get("category_scores", {})
        for category, score in category_scores.items():
            display_category = category.replace("_", " ").title()
            report_lines.append(f"{display_category}: {score * 100:.1f}%")
        
        report_lines.extend([
            "",
            "COMPLIANCE GAPS",
            "===============",
        ])
        
        # Add compliance gaps
        compliance_gaps = results.get("compliance_gaps", [])
        if not compliance_gaps:
            report_lines.append("No specific compliance gaps identified.")
        else:
            for gap in compliance_gaps:
                category = gap.get("category", "").replace("_", " ").title()
                description = gap.get("description", "")
                severity = gap.get("severity", "").upper()
                report_lines.append(f"- {category} ({severity}): {description}")
        
        report_lines.extend([
            "",
            "RECOMMENDATIONS",
            "===============",
        ])
        
        # Add recommendations
        recommendations = results.get("recommendations", [])
        if not recommendations:
            report_lines.append("No specific recommendations provided.")
        else:
            for rec in recommendations:
                category = rec.get("category", "").replace("_", " ").title()
                recommendation = rec.get("recommendation", "")
                priority = rec.get("priority", "").upper()
                report_lines.append(f"- {category} ({priority}): {recommendation}")
        
        # Add prohibition analysis if applicable
        if system_type == "prohibited" and "prohibition_analysis" in results:
            prohibition = results.get("prohibition_analysis", {})
            report_lines.extend([
                "",
                "PROHIBITION ANALYSIS",
                "====================",
                f"Prohibited Category: {prohibition.get('prohibited_category', '').replace('_', ' ').title()}",
                f"Article: {prohibition.get('article', '')}",
                f"Explanation: {prohibition.get('explanation', '')}"
            ])
        
        # Save text report
        text_report_path = self.output_dir / f"{report_filename}.txt"
        with open(text_report_path, "w", encoding="utf-8") as f:
            f.write("\n".join(report_lines))
        
        return text_report_path 